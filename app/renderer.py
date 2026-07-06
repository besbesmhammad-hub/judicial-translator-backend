from __future__ import annotations

import html
import io
import os
import re

import arabic_reshaper
from bidi.algorithm import get_display
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfgen import canvas


def is_arabic(text: str) -> bool:
    return len(re.findall(r"[\u0600-\u06FF]", text)) > len(re.findall(r"[A-Za-z]", text))


def arabic_display(text: str) -> str:
    if not re.search(r"[\u0600-\u06FF]", text or ""):
        return text
    return get_display(arabic_reshaper.reshape(text))


def register_pdf_font() -> str:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/Arial.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont("JudicialSans", path))
                return "JudicialSans"
            except Exception:
                continue
    return "Helvetica"


def split_blocks(text: str) -> list[str]:
    return [block.strip() for block in re.split(r"\n{2,}", text.strip()) if block.strip()]


def parse_markdown_table(block: str) -> list[list[str]] | None:
    lines = [line.strip() for line in block.splitlines() if line.strip()]
    table_lines = [line for line in lines if line.startswith("|") and line.endswith("|")]
    if len(table_lines) < 2:
        return None
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows or None


def clean_marker(value: str) -> tuple[str, str]:
    stripped = value.strip()
    heading = re.match(r"^\[(?:HEADING(?: \d+)?|ARTICLE)\]\s*(.+)$", stripped, re.I | re.S)
    if heading:
        return "heading", heading.group(1).strip()
    page = re.match(r"^\[PAGE \d+\]\s*(.*)$", stripped, re.I | re.S)
    if page:
        return "page", page.group(1).strip()
    return "paragraph", stripped


def render_docx(text: str, title: str = "Translated document") -> bytes:
    document = Document()
    rtl = is_arabic(text)
    heading = document.add_heading(title, level=1)
    if rtl:
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for block in split_blocks(text):
        if block.upper() in {"[TABLE]", "[/TABLE]"}:
            continue
        table_rows = parse_markdown_table(block)
        if table_rows:
            table = document.add_table(rows=len(table_rows), cols=max(len(row) for row in table_rows))
            table.style = "Table Grid"
            for row_index, row in enumerate(table_rows):
                for col_index, cell_value in enumerate(row):
                    cell = table.cell(row_index, col_index)
                    cell.text = cell_value
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT
            continue

        kind, value = clean_marker(block)
        if not value:
            continue
        if kind == "heading":
            paragraph = document.add_heading(value, level=2)
        elif kind == "page":
            paragraph = document.add_paragraph(value)
            paragraph.runs[0].italic = True if paragraph.runs else False
        else:
            paragraph = document.add_paragraph(value)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if rtl else WD_ALIGN_PARAGRAPH.LEFT

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def render_html(text: str, title: str = "Translated document") -> bytes:
    rtl = is_arabic(text)
    direction = "rtl" if rtl else "ltr"
    align = "right" if rtl else "left"
    parts = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        f"<title>{html.escape(title)}</title>",
        "<style>body{font-family:Arial,'Times New Roman',sans-serif;line-height:1.7;margin:36px;} table{border-collapse:collapse;width:100%;margin:16px 0;}td,th{border:1px solid #777;padding:8px;} h1,h2{margin-top:22px;}</style>",
        f"</head><body dir='{direction}' style='text-align:{align}'>",
        f"<h1>{html.escape(title)}</h1>",
    ]
    for block in split_blocks(text):
        table_rows = parse_markdown_table(block)
        if table_rows:
            parts.append("<table>")
            for row in table_rows:
                parts.append("<tr>" + "".join(f"<td>{html.escape(cell)}</td>" for cell in row) + "</tr>")
            parts.append("</table>")
            continue
        kind, value = clean_marker(block)
        if not value:
            continue
        tag = "h2" if kind == "heading" else "p"
        parts.append(f"<{tag}>{html.escape(value).replace(chr(10), '<br>')}</{tag}>")
    parts.append("</body></html>")
    return "".join(parts).encode("utf-8")


def render_pdf(text: str, title: str = "Translated document") -> bytes:
    stream = io.BytesIO()
    pdf = canvas.Canvas(stream, pagesize=A4)
    width, height = A4
    margin = 48
    y = height - margin
    rtl = is_arabic(text)
    font_name = register_pdf_font()
    lines = [title, ""] + text.splitlines()
    font_size = 13 if rtl else 11
    for line in lines:
        if y < margin:
            pdf.showPage()
            y = height - margin
        value = line.strip()
        if not value:
            y -= 12
            continue
        limit = 70 if rtl else 95
        if len(value) > limit:
            chunks = [value[index:index + limit] for index in range(0, len(value), limit)]
        else:
            chunks = [value]
        for chunk in chunks:
            if y < margin:
                pdf.showPage()
                y = height - margin
            pdf.setFont(font_name, font_size)
            if rtl:
                pdf.drawRightString(width - margin, y, arabic_display(chunk))
            else:
                pdf.drawString(margin, y, chunk)
            y -= font_size + 6
    pdf.save()
    return stream.getvalue()


def render_document(text: str, output_format: str, title: str = "Translated document") -> tuple[bytes, str, str]:
    fmt = output_format.lower().strip(".")
    if fmt == "docx":
        return render_docx(text, title), "application/vnd.openxmlformats-officedocument.wordprocessingml.document", "docx"
    if fmt == "pdf":
        return render_pdf(text, title), "application/pdf", "pdf"
    if fmt in {"html", "htm"}:
        return render_html(text, title), "text/html; charset=utf-8", "html"
    if fmt in {"txt", "text"}:
        return text.encode("utf-8"), "text/plain; charset=utf-8", "txt"
    raise ValueError("Unsupported output format")
