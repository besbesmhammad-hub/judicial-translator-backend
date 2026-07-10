from __future__ import annotations

import hashlib
import io
import json
import os
import re
import sys
from pathlib import Path

import fitz
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from app.ocr_utils import choose_better_text, is_low_quality_text, normalize_text, ocr_pdf_page_text


CORPUS_PATH = ROOT / "app" / "data" / "tunisian_legal_corpus.jsonl"
DOWNLOADS = Path.home() / "Downloads"


DOCS = [
    {
        "filename": "compte-rendu-stagiaire-v2025.pdf",
        "doc_id": "formulaire_compte_rendu_stagiaire",
        "title": "Formulaire de compte rendu de stagiaire CCT",
        "authority": "Compagnie des Comptables de Tunisie",
        "source_tier": "form_template",
        "year": 2012,
        "domain": "stage_professionnel",
    },
    {
        "filename": "lettre-circulaire-stagiaires.pdf",
        "doc_id": "circulaire_stagiaires_2018",
        "title": "Lettre circulaire stagiaires 2018",
        "authority": "Compagnie des Comptables de Tunisie",
        "source_tier": "professional_circular",
        "year": 2018,
        "domain": "stage_professionnel",
    },
    {
        "filename": None,
        "doc_id": "textes_profession_comptable_2018",
        "title": "Textes relatifs aux comptables, experts-comptables et commissaires aux comptes",
        "authority": "Imprimerie Officielle de la Republique Tunisienne",
        "source_tier": "professional_text_collection",
        "year": 2018,
        "domain": "reglementation_professionnelle",
    },
    {
        "filename": "Rapport-moral-2023.pdf",
        "doc_id": "rapport_moral_2023",
        "title": "Rapport moral 2023",
        "authority": "Organisation professionnelle comptable tunisienne",
        "source_tier": "institutional_report",
        "year": 2023,
        "domain": "vie_professionnelle",
    },
    {
        "filename": "rapport-moral-2024.pdf",
        "doc_id": "rapport_moral_2024",
        "title": "Rapport moral 2024",
        "authority": "Organisation professionnelle comptable tunisienne",
        "source_tier": "institutional_report",
        "year": 2024,
        "domain": "vie_professionnelle",
    },
    {
        "filename": "rapport-moral-2025.pdf",
        "doc_id": "rapport_moral_2025",
        "title": "Rapport moral 2025",
        "authority": "Organisation professionnelle comptable tunisienne",
        "source_tier": "institutional_report",
        "year": 2025,
        "domain": "vie_professionnelle",
    },
    {
        "filename": "CODE DE COMMERCE.pdf",
        "doc_id": "code_commerce_2014",
        "title": "Code de commerce tunisien",
        "authority": "Imprimerie Officielle de la Republique Tunisienne",
        "source_tier": "primary_law",
        "year": 2014,
        "domain": "droit_commercial",
    },
    {
        "filename": "CODE COC.pdf",
        "doc_id": "code_obligations_contrats_2015",
        "title": "Code des obligations et des contrats",
        "authority": "Imprimerie Officielle de la Republique Tunisienne",
        "source_tier": "primary_law",
        "year": 2015,
        "domain": "droit_civil_commercial",
    },
    {
        "filename": "code_societes_fr.pdf",
        "doc_id": "code_societes_commerciales_2022",
        "title": "Code des societes commerciales",
        "authority": "Imprimerie Officielle de la Republique Tunisienne",
        "source_tier": "primary_law",
        "year": 2022,
        "domain": "droit_societes",
    },
    {
        "filename": "Guide-inscription-personnes-morales-societes-2026.pdf",
        "doc_id": "guide_inscription_personnes_morales_2026",
        "title": "Guide inscription personnes morales et societes 2026",
        "authority": "Compagnie des Comptables de Tunisie",
        "source_tier": "professional_guide",
        "year": 2026,
        "domain": "reglementation_professionnelle",
    },
    {
        "filename": "Guide-inscription-stagiaires-2026.pdf",
        "doc_id": "guide_inscription_stagiaires_2026",
        "title": "Guide inscription stagiaires 2026",
        "authority": "Compagnie des Comptables de Tunisie",
        "source_tier": "professional_guide",
        "year": 2026,
        "domain": "stage_professionnel",
    },
    {
        "filename": "Guide-inscription-personnes-physiques-2026.pdf",
        "doc_id": "guide_inscription_personnes_physiques_2026",
        "title": "Guide inscription personnes physiques 2026",
        "authority": "Compagnie des Comptables de Tunisie",
        "source_tier": "professional_guide",
        "year": 2026,
        "domain": "reglementation_professionnelle",
    },
    {
        "filename": "Formulaire-Demande-De-Radiation-2026.pdf",
        "doc_id": "formulaire_radiation_2026",
        "title": "Formulaire demande de radiation 2026",
        "authority": "Compagnie des Comptables de Tunisie",
        "source_tier": "form_template",
        "year": 2026,
        "domain": "reglementation_professionnelle",
    },
    {
        "filename": "Demande-dune-attestation-dinscription-2026.pdf",
        "doc_id": "demande_attestation_inscription_2026",
        "title": "Demande d attestation d inscription 2026",
        "authority": "Instance professionnelle comptable tunisienne",
        "source_tier": "form_template",
        "year": 2026,
        "domain": "reglementation_professionnelle",
    },
    {
        "filename": "Formulaire-Demande-De-Suspension-2026.pdf",
        "doc_id": "formulaire_suspension_2026",
        "title": "Formulaire demande de suspension 2026",
        "authority": "Compagnie des Comptables de Tunisie",
        "source_tier": "form_template",
        "year": 2026,
        "domain": "reglementation_professionnelle",
    },
    {
        "filename": "Les tribunaux de premières instances.docx",
        "doc_id": "tribunaux_premiere_instance_guide",
        "title": "Les tribunaux de premiere instance",
        "authority": "Source explicative judiciaire tunisienne",
        "source_tier": "secondary_legal_guide",
        "year": 2026,
        "domain": "organisation_judiciaire",
    },
    {
        "filename": "La Cour de Cassation.docx",
        "doc_id": "cour_cassation_guide",
        "title": "La Cour de cassation",
        "authority": "Source explicative judiciaire tunisienne",
        "source_tier": "secondary_legal_guide",
        "year": 2026,
        "domain": "organisation_judiciaire",
    },
    {
        "filename": "sa.pdf",
        "doc_id": "checklist_constitution_sa_api",
        "title": "Checklist constitution d une societe anonyme",
        "authority": "Agence de Promotion de l Industrie",
        "source_tier": "administrative_checklist",
        "year": 2026,
        "domain": "creation_societe",
    },
    {
        "filename": "SARL Tunisie.docx",
        "doc_id": "guide_creation_sarl_tunisie",
        "title": "Guide creation SARL Tunisie",
        "authority": "Source explicative pratique tunisienne",
        "source_tier": "secondary_legal_guide",
        "year": 2021,
        "domain": "creation_societe",
    },
    {
        "filename": "La fermeture d.docx",
        "doc_id": "guide_fermeture_entreprise_tunisie",
        "title": "Guide fermeture d entreprise en Tunisie",
        "authority": "Source explicative pratique tunisienne",
        "source_tier": "secondary_legal_guide",
        "year": 2026,
        "domain": "dissolution_liquidation",
    },
    {
        "filename": "Tunisie.docx",
        "doc_id": "cassation_chambres_reunies_terrorisme_2019",
        "title": "Cour de cassation chambres reunies 11 avril 2019",
        "authority": "Jurisprudence tunisienne",
        "source_tier": "case_law",
        "year": 2019,
        "domain": "procedure_penale_terrorisme",
    },
    {
        "filename": "Tunisie1.docx",
        "doc_id": "cassation_acte_commerce_accessoire_2019",
        "title": "Cour de cassation acte de commerce par accessoire 4 mars 2019",
        "authority": "Jurisprudence tunisienne",
        "source_tier": "case_law",
        "year": 2019,
        "domain": "droit_commercial_jurisprudence",
    },
    {
        "filename": "Tunisie2.docx",
        "doc_id": "cassation_sequestre_societe_anonyme_2018",
        "title": "Cour de cassation sequestre societe anonyme 29 mai 2018",
        "authority": "Jurisprudence tunisienne",
        "source_tier": "case_law",
        "year": 2018,
        "domain": "droit_societes_jurisprudence",
    },
    {
        "filename": "Tunisie3.docx",
        "doc_id": "cassation_arbitrage_interne_2018",
        "title": "Cour de cassation arbitrage interne 26 avril 2018",
        "authority": "Jurisprudence tunisienne",
        "source_tier": "case_law",
        "year": 2018,
        "domain": "arbitrage_jurisprudence",
    },
    {
        "filename": "nistie efface le délit ainsi que la sanction.docx",
        "doc_id": "analyse_amnistie_reconciliation_administrative",
        "title": "Amnistie et reconciliation nationale administrative",
        "authority": "Analyse juridique tunisienne",
        "source_tier": "jurisprudence_analysis",
        "year": 2017,
        "domain": "droit_penal_administratif",
    },
    {
        "filename": "Tunisie4.docx",
        "doc_id": "cassation_dissolution_sarl_affectio_2018",
        "title": "Cour de cassation dissolution SARL affectio societatis 26 mars 2018",
        "authority": "Jurisprudence tunisienne",
        "source_tier": "case_law",
        "year": 2018,
        "domain": "droit_societes_jurisprudence",
    },
    {
        "filename": "Tunisie5.docx",
        "doc_id": "cassation_reglement_judiciaire_cotisations_2017",
        "title": "Cour de cassation reglement judiciaire et cotisations 8 decembre 2017",
        "authority": "Jurisprudence tunisienne",
        "source_tier": "case_law",
        "year": 2017,
        "domain": "entreprises_en_difficulte_jurisprudence",
    },
    {
        "filename": "Tunisie6.docx",
        "doc_id": "cassation_terrorisme_participation_groupe_2017",
        "title": "Cour de cassation participation a un groupe terroriste 30 novembre 2017",
        "authority": "Jurisprudence tunisienne",
        "source_tier": "case_law",
        "year": 2017,
        "domain": "procedure_penale_terrorisme",
    },
    {
        "filename": "Tunisie7.docx",
        "doc_id": "cassation_accident_route_baremes_2017",
        "title": "Cour de cassation accidents de la voie publique et baremes 9 novembre 2017",
        "authority": "Jurisprudence tunisienne",
        "source_tier": "case_law",
        "year": 2017,
        "domain": "assurance_responsabilite_jurisprudence",
    },
    {
        "filename": "Tunisie8.docx",
        "doc_id": "cassation_clause_compromissoire_2017",
        "title": "Cour de cassation clause compromissoire et promesse de vente 6 novembre 2017",
        "authority": "Jurisprudence tunisienne",
        "source_tier": "case_law",
        "year": 2017,
        "domain": "arbitrage_jurisprudence",
    },
    {
        "filename": "bo170411.pdf",
        "doc_id": "cmf_bulletin_officiel_2017_04_11",
        "title": "Bulletin Officiel du Conseil du Marche Financier 11 avril 2017",
        "authority": "Conseil du Marche Financier",
        "source_tier": "regulatory_bulletin",
        "year": 2017,
        "domain": "marche_financier_reglementation",
        "max_pages": 30,
    },
    {
        "filename": "1700485732.pdf",
        "doc_id": "guide_agrement_etablissement_paiement_tunisie",
        "title": "Guide d agrement pour la creation d un etablissement de paiement en Tunisie",
        "authority": "Guide reglementaire tunisien",
        "source_tier": "regulatory_guidance",
        "year": 2023,
        "domain": "paiement_reglementation",
        "max_pages": 24,
    },
    {
        "filename": "e155fd21-5b8f-47cb-b8e9-14e1cf16939b_1772008340_51eb70b554dd29c7446b070ae1fa5167.pdf",
        "doc_id": "appel_offres_assurance_tunisie_autoroutes_2026",
        "title": "Appel d offres assurance Tunisie Autoroutes 2026",
        "authority": "Societe Tunisie Autoroutes / Ministere de l Equipement et de l Habitat",
        "source_tier": "public_procurement_tender",
        "year": 2026,
        "domain": "commande_publique_assurance",
        "max_pages": 20,
    },
    {
        "filename": "xabyt-prospectus-06-01-2012.pdf",
        "doc_id": "prospectus_hexabyte_2011_2012",
        "title": "Prospectus Hexabyte introduction au marche alternatif",
        "authority": "Conseil du Marche Financier",
        "source_tier": "market_prospectus",
        "year": 2011,
        "domain": "marche_financier_prospectus",
        "max_pages": 20,
    },
    {
        "filename": "Prospectus_fusion_tunisie_Leasing.pdf",
        "doc_id": "prospectus_fusion_tunisie_leasing",
        "title": "Prospectus fusion Tunisie Leasing",
        "authority": "Conseil du Marche Financier",
        "source_tier": "market_prospectus",
        "year": 2011,
        "domain": "marche_financier_prospectus",
        "max_pages": 24,
    },
    {
        "filename": "Strategie-Habitat-Tunisie_20150427_final.pdf",
        "doc_id": "strategie_habitat_tunisie_2015",
        "title": "Vers une nouvelle strategie de l habitat en Tunisie",
        "authority": "Ministere de l Equipement, de l Amenagement du Territoire et du Developpement Durable",
        "source_tier": "policy_strategy",
        "year": 2015,
        "domain": "politique_publique_habitat",
        "max_pages": 18,
    },
    {
        "filename": "multi-page.pdf",
        "doc_id": "banque_mondiale_strategie_transports_tunisie",
        "title": "Banque mondiale etude sur la strategie des transports Tunisie",
        "authority": "Banque mondiale",
        "source_tier": "external_report",
        "year": 0,
        "domain": "transport_politique_publique",
        "max_pages": 18,
    },
    {
        "filename": "BILAN 2016.pdf",
        "doc_id": "rapport_cac_ance_2016",
        "title": "Rapports de revision des comptes ANCE 2016",
        "authority": "Cabinet d expertise comptable / commissariat aux comptes",
        "source_tier": "audit_report",
        "year": 2016,
        "domain": "audit_comptable",
        "max_pages": 14,
    },
    {
        "filename": "INNORPI RAPPORTS CAC 2021.pdf",
        "doc_id": "rapport_cac_innorpi_2021",
        "title": "Rapports du reviseur des comptes INNORPI 2021",
        "authority": "Cabinet d expertise comptable / commissariat aux comptes",
        "source_tier": "audit_report",
        "year": 2021,
        "domain": "audit_comptable",
        "max_pages": 14,
    },
    {
        "filename": "Rapport-general-et-special-2018-2.pdf",
        "doc_id": "rapport_cac_bna_2018",
        "title": "Rapport general et special BNA 2018",
        "authority": "Cabinets d expertise comptable / commissariat aux comptes",
        "source_tier": "audit_report",
        "year": 2018,
        "domain": "audit_bancaire",
        "max_pages": 12,
    },
    {
        "filename": "rapport_cac_ote_2014.pdf",
        "doc_id": "rapport_cac_ote_2014",
        "title": "Rapport general du commissaire aux comptes OTE 2014",
        "authority": "Cabinet d expertise comptable / commissariat aux comptes",
        "source_tier": "audit_report",
        "year": 2014,
        "domain": "audit_comptable",
        "max_pages": 10,
    },
    {
        "filename": "RAPPORT-CAC-Cefa-Tunisie-exercice-2020.pdf",
        "doc_id": "rapport_cac_cefa_tunisie_2020",
        "title": "Rapport du commissaire aux comptes CEFA Tunisie 2020",
        "authority": "Cabinet d expertise comptable / commissariat aux comptes",
        "source_tier": "audit_report",
        "year": 2020,
        "domain": "audit_comptable",
        "max_pages": 12,
    },
    {
        "filename": "cac-rapport-act-2021.pdf",
        "doc_id": "rapport_cac_act_2021",
        "title": "Rapport CAC ACT 2021",
        "authority": "Cabinet d expertise comptable / commissariat aux comptes",
        "source_tier": "audit_report",
        "year": 2021,
        "domain": "audit_association",
        "max_pages": 10,
    },
    {
        "filename": "IRC - Rapport CAC 2017.pdf",
        "doc_id": "rapport_cac_irc_2017",
        "title": "Rapport d audit IRC Tunisie 2017",
        "authority": "Cabinet d expertise comptable / commissariat aux comptes",
        "source_tier": "audit_report",
        "year": 2017,
        "domain": "audit_association",
        "max_pages": 12,
    },
    {
        "filename": "note_d_orientation_circulaire_bct_2012_02.pdf",
        "doc_id": "note_orientation_bct_2012_02",
        "title": "Note d orientation sur la circulaire BCT 2012-02",
        "authority": "Ordre des Experts Comptables de Tunisie",
        "source_tier": "professional_guidance",
        "year": 2012,
        "domain": "audit_bancaire_guidance",
        "max_pages": 14,
    },
    {
        "filename": "Rapport réviseur légal_EF_2017.pdf",
        "doc_id": "rapport_reviseur_legal_smls_2017",
        "title": "Rapport reviseur legal et etats financiers 2017",
        "authority": "Cabinet d expertise comptable / commissariat aux comptes",
        "source_tier": "audit_report",
        "year": 2017,
        "domain": "audit_comptable",
        "max_pages": 12,
    },
    {
        "filename": "RAPPORT_GENERAL_DES_CAC_2017_C.pdf",
        "doc_id": "rapport_general_cac_2017",
        "title": "Rapport general des commissaires aux comptes 2017",
        "authority": "Cabinets d expertise comptable / commissariat aux comptes",
        "source_tier": "audit_report",
        "year": 2017,
        "domain": "audit_comptable",
        "max_pages": 10,
    },
    {
        "filename": "Rapport-daudit-2023.pdf",
        "doc_id": "rapport_audit_nebras_2023",
        "title": "Rapport du commissaire aux comptes Nebras 2023",
        "authority": "Cabinet d expertise comptable / commissariat aux comptes",
        "source_tier": "audit_report",
        "year": 2023,
        "domain": "audit_association",
        "max_pages": 12,
    },
    {
        "filename": "L_1993_61fr.pdf",
        "doc_id": "loi_experts_judiciaires_1993",
        "title": "Loi relative aux experts judiciaires 23 juin 1993",
        "authority": "Journal Officiel de la Republique Tunisienne",
        "source_tier": "primary_law",
        "year": 1993,
        "domain": "expertise_judiciaire",
        "max_pages": 3,
    },
    {
        "filename": "Arrete_11_8_93_fr_compostion.pdf",
        "doc_id": "arrete_composition_commission_experts_1993",
        "title": "Arrete fixant la composition de la commission regionale des experts judiciaires",
        "authority": "Ministere de la Justice",
        "source_tier": "implementing_regulation",
        "year": 1993,
        "domain": "expertise_judiciaire",
        "max_pages": 1,
    },
    {
        "filename": "A_11aout93fr.pdf",
        "doc_id": "arrete_delais_inscription_experts_1993",
        "title": "Arrete fixant les delais de presentation des demandes d inscription des experts judiciaires",
        "authority": "Ministere de la Justice",
        "source_tier": "implementing_regulation",
        "year": 1993,
        "domain": "expertise_judiciaire",
        "max_pages": 1,
    },
    {
        "filename": "Arrete_03_6_2000_fr.pdf",
        "doc_id": "arrete_manuel_procedures_expert_judiciaire_2000",
        "title": "Arrete relatif a l approbation du manuel de procedures de l expert judiciaire",
        "authority": "Ministere de la Justice",
        "source_tier": "implementing_regulation",
        "year": 2000,
        "domain": "expertise_judiciaire",
        "max_pages": 1,
    },
    {
        "filename": "loi_33-2010_fr.pdf",
        "doc_id": "loi_modification_experts_judiciaires_2010",
        "title": "Loi modifiant le regime des experts judiciaires 2010",
        "authority": "Journal Officiel de la Republique Tunisienne",
        "source_tier": "primary_law",
        "year": 2010,
        "domain": "expertise_judiciaire",
        "max_pages": 2,
    },
    {
        "filename": "fich_091_035.pdf",
        "doc_id": "article_revue_expertise_comptable_2011",
        "title": "Article Revue Comptable et Financiere expertise comptable 2011",
        "authority": "Revue Comptable et Financiere",
        "source_tier": "professional_article",
        "year": 2011,
        "domain": "profession_expertise_comptable",
        "max_pages": 8,
    },
    {
        "filename": "cadre-conceptuel-de-l_information-financic3a8re.pdf",
        "doc_id": "ifrs_cadre_conceptuel_information_financiere",
        "title": "Cadre conceptuel de l information financiere",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2010,
        "domain": "ifrs_cadre_conceptuel",
        "max_pages": 18,
    },
    {
        "filename": "ifrs-1-premic3a8re-application-des-normes-internationales-d_information-financic3a8re.pdf",
        "doc_id": "ifrs_1_premiere_application",
        "title": "IFRS 1 Premiere application des normes internationales d information financiere",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2009,
        "domain": "ifrs_premiere_application",
        "max_pages": 16,
    },
    {
        "filename": "ifrs-2-paiement-fondc3a9-sur-des-actions.pdf",
        "doc_id": "ifrs_2_paiement_fonde_sur_actions",
        "title": "IFRS 2 Paiement fonde sur des actions",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2004,
        "domain": "ifrs_paiement_actions",
        "max_pages": 16,
    },
    {
        "filename": "ifrs-3-regroupements-d_entreprises.pdf",
        "doc_id": "ifrs_3_regroupements_entreprises",
        "title": "IFRS 3 Regroupements d entreprises",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2008,
        "domain": "ifrs_regroupements_entreprises",
        "max_pages": 18,
    },
    {
        "filename": "ifrs-4-contrats-d_assurance.pdf",
        "doc_id": "ifrs_4_contrats_assurance",
        "title": "IFRS 4 Contrats d assurance",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2004,
        "domain": "ifrs_assurance",
        "max_pages": 14,
    },
    {
        "filename": "ifrs-5-actifs-non-courants-dc3a9tenus-en-vue-de-la-vente-et-activitc3a9s-abandonnc3a9es.pdf",
        "doc_id": "ifrs_5_actifs_non_courants_vente",
        "title": "IFRS 5 Actifs non courants detenus en vue de la vente et activites abandonnees",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2004,
        "domain": "ifrs_actifs_non_courants_vente",
        "max_pages": 11,
    },
    {
        "filename": "ifrs-6-prospection-et-c3a9valuation-de-ressources-minc3a9rales.pdf",
        "doc_id": "ifrs_6_ressources_minieres",
        "title": "IFRS 6 Prospection et evaluation de ressources minerales",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2004,
        "domain": "ifrs_ressources_minieres",
        "max_pages": 5,
    },
    {
        "filename": "ifrs-7-instruments-financiers-informations-c3a0-fournir.pdf",
        "doc_id": "ifrs_7_instruments_financiers_informations",
        "title": "IFRS 7 Instruments financiers Informations a fournir",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2005,
        "domain": "ifrs_information_instruments_financiers",
        "max_pages": 16,
    },
    {
        "filename": "ifrs-8-secteurs-opc3a9rationnels.pdf",
        "doc_id": "ifrs_8_secteurs_operationnels",
        "title": "IFRS 8 Secteurs operationnels",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2006,
        "domain": "ifrs_secteurs_operationnels",
        "max_pages": 9,
    },
    {
        "filename": "ifrs-09.pdf",
        "doc_id": "ifrs_9_instruments_financiers",
        "title": "IFRS 9 Instruments financiers",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2014,
        "domain": "ifrs_instruments_financiers",
        "max_pages": 20,
    },
    {
        "filename": "ifrs-10-c3a9tats-financiers-consolidc3a9s.pdf",
        "doc_id": "ifrs_10_etats_financiers_consolides",
        "title": "IFRS 10 Etats financiers consolides",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2011,
        "domain": "ifrs_consolidation",
        "max_pages": 18,
    },
    {
        "filename": "ifrs-11-partenariats.pdf",
        "doc_id": "ifrs_11_partenariats",
        "title": "IFRS 11 Partenariats",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2011,
        "domain": "ifrs_partenariats",
        "max_pages": 14,
    },
    {
        "filename": "ifrs-12-informations-c3a0-fournir-sur-les-intc3a9rc3aats-dc3a9tenus-dans-d_autres-entitc3a9s.pdf",
        "doc_id": "ifrs_12_interets_autres_entites",
        "title": "IFRS 12 Informations a fournir sur les interets detenus dans d autres entites",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2011,
        "domain": "ifrs_information_autres_entites",
        "max_pages": 12,
    },
    {
        "filename": "ifrs-13-c3a9valuation-de-la-juste-valeur.pdf",
        "doc_id": "ifrs_13_juste_valeur",
        "title": "IFRS 13 Evaluation de la juste valeur",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2011,
        "domain": "ifrs_juste_valeur",
        "max_pages": 16,
    },
    {
        "filename": "ifrs14.pdf",
        "doc_id": "ifrs_14_comptes_report_reglementaires",
        "title": "IFRS 14 Comptes de report reglementaires",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2014,
        "domain": "ifrs_report_reglementaire",
        "max_pages": 10,
    },
    {
        "filename": "ifrs15.pdf",
        "doc_id": "ifrs_15_produits_contrats_clients",
        "title": "IFRS 15 Produits des activites ordinaires tires de contrats conclus avec des clients",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2014,
        "domain": "ifrs_revenus_clients",
        "max_pages": 18,
    },
    {
        "filename": "ifrs16.pdf",
        "doc_id": "ifrs_16_contrats_location",
        "title": "IFRS 16 Contrats de location",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2016,
        "domain": "ifrs_location",
        "max_pages": 18,
    },
    {
        "filename": "ias-1-prc3a9sentation-des-c3a9tats-financiers.pdf",
        "doc_id": "ias_1_presentation_etats_financiers",
        "title": "IAS 1 Presentation des etats financiers",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2007,
        "domain": "ias_presentation_etats_financiers",
        "max_pages": 14,
    },
    {
        "filename": "ias-2-stocks.pdf",
        "doc_id": "ias_2_stocks",
        "title": "IAS 2 Stocks",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2003,
        "domain": "ias_stocks",
        "max_pages": 6,
    },
    {
        "filename": "ias-7-tableau-des-flux-de-trc3a9sorerie.pdf",
        "doc_id": "ias_7_tableau_flux_tresorerie",
        "title": "IAS 7 Tableau des flux de tresorerie",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2007,
        "domain": "ias_flux_tresorerie",
        "max_pages": 8,
    },
    {
        "filename": "ias-8-mc3a9thodes-comptables-changements-d_estimations-comptables-et-erreurs.pdf",
        "doc_id": "ias_8_methodes_comptables_estimations_erreurs",
        "title": "IAS 8 Methodes comptables changements d estimations comptables et erreurs",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2005,
        "domain": "ias_methodes_estimations_erreurs",
        "max_pages": 8,
    },
    {
        "filename": "ias-10-c3a9vc3a9nements-postc3a9rieurs-c3a0-la-date-de-clc3b4ture.pdf",
        "doc_id": "ias_10_evenements_post_cloture",
        "title": "IAS 10 Evenements posterieurs a la date de cloture",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2003,
        "domain": "ias_evenements_post_cloture",
        "max_pages": 4,
    },
    {
        "filename": "ias-11-contrats-de-construction.pdf",
        "doc_id": "ias_11_contrats_construction",
        "title": "IAS 11 Contrats de construction",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 1993,
        "domain": "ias_contrats_construction",
        "max_pages": 7,
    },
    {
        "filename": "ias-12-impc3b4ts-sur-le-rc3a9sultat.pdf",
        "doc_id": "ias_12_impots_resultat",
        "title": "IAS 12 Impots sur le resultat",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 1998,
        "domain": "ias_impots_resultat",
        "max_pages": 14,
    },
    {
        "filename": "ias-16-immobilisations-corporelles.pdf",
        "doc_id": "ias_16_immobilisations_corporelles",
        "title": "IAS 16 Immobilisations corporelles",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2003,
        "domain": "ias_immobilisations_corporelles",
        "max_pages": 10,
    },
    {
        "filename": "ias-17-contrats-de-location.pdf",
        "doc_id": "ias_17_contrats_location",
        "title": "IAS 17 Contrats de location",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2003,
        "domain": "ias_location",
        "max_pages": 9,
    },
    {
        "filename": "ias-18-produits-des-activitc3a9s-ordinaires.pdf",
        "doc_id": "ias_18_produits_activites_ordinaires",
        "title": "IAS 18 Produits des activites ordinaires",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 1993,
        "domain": "ias_revenus",
        "max_pages": 6,
    },
    {
        "filename": "ias-19-avantages-du-personnel.pdf",
        "doc_id": "ias_19_avantages_personnel",
        "title": "IAS 19 Avantages du personnel",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2011,
        "domain": "ias_avantages_personnel",
        "max_pages": 14,
    },
    {
        "filename": "ias-20-comptabilisation-des-subventions-publiques-et-informations-c3a0-fournir-sur-l_aide-publique.pdf",
        "doc_id": "ias_20_subventions_publiques_aide_publique",
        "title": "IAS 20 Comptabilisation des subventions publiques et informations a fournir sur l aide publique",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 1983,
        "domain": "ias_subventions_publiques",
        "max_pages": 5,
    },
    {
        "filename": "ias-21-effets-des-variations-des-cours-des-monnaies-c3a9trangc3a8res.pdf",
        "doc_id": "ias_21_variations_cours_monnaies_etrangeres",
        "title": "IAS 21 Effets des variations des cours des monnaies etrangeres",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2003,
        "domain": "ias_monnaies_etrangeres",
        "max_pages": 9,
    },
    {
        "filename": "ias-23-coc3bbts-d_emprunt.pdf",
        "doc_id": "ias_23_couts_emprunt",
        "title": "IAS 23 Couts d emprunt",
        "authority": "IFRS Foundation / IASB",
        "source_tier": "accounting_standard",
        "year": 2007,
        "domain": "ias_couts_emprunt",
        "max_pages": 4,
    },
]

DOC_ID_FILTER = {value.strip() for value in os.environ.get("DOC_IDS", "").split(",") if value.strip()}
MAX_PAGES_PER_DOC = int(os.environ.get("MAX_PAGES_PER_DOC", "0") or "0")
START_PAGE = int(os.environ.get("START_PAGE", "1") or "1")
END_PAGE = int(os.environ.get("END_PAGE", "0") or "0")


HEADING_RE = re.compile(
    r"^(article\s+\d+|art\.\s*\d+|chapitre\s+[ivxlcdm\d]+|section\s+\d+|titre\s+[ivxlcdm\d]+|annexe|appendice|partie\s+[ivxlcdm\d]+|i\.|ii\.|iii\.|iv\.|v\.|vi\.|vii\.|viii\.|ix\.|x\.|\d+\s*[\.)-])",
    re.I,
)


def resolve_filename(meta: dict) -> str:
    if meta["filename"]:
        return meta["filename"]
    for path in DOWNLOADS.iterdir():
        if path.suffix.lower() == ".pdf" and not path.name.isascii() and path.stat().st_size == 1_196_628:
            return path.name
    raise FileNotFoundError("Arabic professional text collection PDF not found.")


def pick_heading(paragraph: str) -> str:
    for raw_line in paragraph.splitlines()[:4]:
        line = raw_line.strip(" -:•\t")
        if not line:
            continue
        if len(line) <= 140 and (HEADING_RE.match(line) or line.isupper()):
            return line
    return ""


def chunk_text(text: str, limit: int = 1800) -> list[tuple[str, str]]:
    paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    chunks: list[tuple[str, str]] = []
    current = ""
    current_heading = ""
    for paragraph in paragraphs:
        paragraph_heading = pick_heading(paragraph)
        candidate = paragraph if not current else f"{current}\n\n{paragraph}"
        if current and len(candidate) > limit:
            chunks.append((current_heading, current.strip()))
            current = paragraph
            current_heading = paragraph_heading or current_heading or ""
        else:
            current = candidate
            current_heading = current_heading or paragraph_heading or ""
        if len(current) > limit * 1.35:
            parts = [current[i:i + limit] for i in range(0, len(current), limit)]
            for part in parts[:-1]:
                chunks.append((current_heading, part.strip()))
            current = parts[-1].strip()
    if current.strip():
        chunks.append((current_heading, current.strip()))
    return chunks


def extract_page_text(page) -> str:
    blocks = page.get_text("blocks")
    embedded = normalize_text("\n".join(str(block[4]).strip() for block in blocks if len(block) > 4))
    if is_low_quality_text(embedded):
        try:
            return choose_better_text(embedded, ocr_pdf_page_text(page))
        except Exception:
            return embedded
    return embedded


def extract_docx_text(path: Path) -> str:
    document = Document(path)
    blocks: list[str] = []
    for paragraph in document.paragraphs:
        value = normalize_text(paragraph.text.strip())
        if not value:
            continue
        style = (paragraph.style.name or "").lower() if paragraph.style else ""
        if "heading" in style or "titre" in style or "title" in style:
            blocks.append(f"[HEADING] {value}")
        elif re.match(r"^(article|clause|section|chapitre)\s+", value, re.I):
            blocks.append(f"[ARTICLE] {value}")
        else:
            blocks.append(value)

    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [normalize_text(cell.text.replace("\n", " ").strip().replace("|", "/")) for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            blocks.extend(
                ["[TABLE]", "| " + " | ".join(rows[0]) + " |", "| " + " | ".join(["---"] * width) + " |"]
                + ["| " + " | ".join(row) + " |" for row in rows[1:]]
                + ["[/TABLE]"]
            )
    return "\n\n".join(blocks).strip()


def build_records(meta: dict) -> list[dict]:
    path = DOWNLOADS / resolve_filename(meta)
    records: list[dict] = []
    if path.suffix.lower() == ".docx":
        text = extract_docx_text(path)
        if not text:
            return records
        for local_index, (heading, chunk) in enumerate(chunk_text(text), start=1):
            digest = hashlib.blake2b(
                f"{meta['doc_id']}|1|{local_index}|{chunk[:200]}".encode("utf-8"),
                digest_size=8,
            ).hexdigest()
            records.append({
                "id": digest,
                "doc_id": meta["doc_id"],
                "title": meta["title"],
                "filename": path.name,
                "page": 1,
                "heading": heading,
                "text": chunk,
                "authority": meta["authority"],
                "source_tier": meta["source_tier"],
                "year": meta["year"],
                "domain": meta["domain"],
            })
        return records

    doc = fitz.open(path)
    start_index = max(0, START_PAGE - 1)
    end_index = doc.page_count
    doc_page_cap = int(meta.get("max_pages", 0) or 0)
    if END_PAGE:
        end_index = min(doc.page_count, END_PAGE)
    elif MAX_PAGES_PER_DOC:
        end_index = min(doc.page_count, start_index + MAX_PAGES_PER_DOC)
    elif doc_page_cap:
        end_index = min(doc.page_count, start_index + doc_page_cap)
    for page_index in range(start_index, end_index):
        page_text = extract_page_text(doc.load_page(page_index))
        if not page_text:
            continue
        for local_index, (heading, chunk) in enumerate(chunk_text(page_text), start=1):
            digest = hashlib.blake2b(
                f"{meta['doc_id']}|{page_index + 1}|{local_index}|{chunk[:200]}".encode("utf-8"),
                digest_size=8,
            ).hexdigest()
            records.append({
                "id": digest,
                "doc_id": meta["doc_id"],
                "title": meta["title"],
                "filename": path.name,
                "page": page_index + 1,
                "heading": heading,
                "text": chunk,
                "authority": meta["authority"],
                "source_tier": meta["source_tier"],
                "year": meta["year"],
                "domain": meta["domain"],
            })
    return records


def main() -> None:
    selected_docs = [meta for meta in DOCS if not DOC_ID_FILTER or meta["doc_id"] in DOC_ID_FILTER]
    existing = []
    with CORPUS_PATH.open("r", encoding="utf-8") as handle:
        for line in handle:
            if line.strip():
                existing.append(json.loads(line))

    target_ids = {meta["doc_id"] for meta in selected_docs}
    page_range_mode = bool(END_PAGE or START_PAGE > 1 or MAX_PAGES_PER_DOC)

    if page_range_mode:
        def keep_row(row: dict) -> bool:
            if row.get("doc_id") not in target_ids:
                return True
            page = int(row.get("page") or 0)
            range_end = END_PAGE if END_PAGE else START_PAGE + MAX_PAGES_PER_DOC - 1 if MAX_PAGES_PER_DOC else 10**9
            return not (START_PAGE <= page <= range_end)
        kept = [row for row in existing if keep_row(row)]
    else:
        kept = [row for row in existing if row.get("doc_id") not in target_ids]
    rebuilt: list[dict] = []
    for meta in selected_docs:
        rebuilt.extend(build_records(meta))

    with CORPUS_PATH.open("w", encoding="utf-8") as handle:
        for row in kept + rebuilt:
            handle.write(json.dumps(row, ensure_ascii=False, separators=(",", ":")) + "\n")

    print(f"kept={len(kept)} rebuilt={len(rebuilt)} total={len(kept) + len(rebuilt)}")
    for doc_id in sorted(target_ids):
        print(doc_id, sum(1 for row in rebuilt if row["doc_id"] == doc_id))


if __name__ == "__main__":
    main()
